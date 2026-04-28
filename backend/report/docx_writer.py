"""DOCX report generator using python-docx."""

from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_report(
    filepath: str,
    topic: str,
    sub_questions: list[str],
    research_findings: dict[str, str],
    reflection_gaps: list[str],
    final_synthesis: str,
    confidence: str,
    confidence_score: float,
    web_used: bool,
) -> None:
    """Generate a 7-section research report as DOCX."""
    doc = Document()

    # ─── Styles ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ─── 1. Title Page ───────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DocMind AI Research Report")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(91, 108, 246)
    run.bold = True

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(topic)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    meta.add_run(f"Generated on {date_str}\n").font.size = Pt(10)
    conf_run = meta.add_run(f"Confidence: {confidence.upper()} ({confidence_score:.2%})")
    conf_run.font.size = Pt(10)
    conf_colors = {"high": RGBColor(34, 139, 34), "medium": RGBColor(218, 165, 32), "low": RGBColor(220, 20, 60)}
    conf_run.font.color.rgb = conf_colors.get(confidence, RGBColor(128, 128, 128))

    doc.add_page_break()

    # ─── 2. Executive Summary ────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    # Extract first 2-3 paragraphs from synthesis
    synth_paragraphs = [p.strip() for p in final_synthesis.split("\n\n") if p.strip()]
    summary = synth_paragraphs[:3] if synth_paragraphs else ["No summary available."]
    for para in summary:
        doc.add_paragraph(para)

    # ─── 3. Research Sub-Questions ───────────────────────────
    doc.add_heading("Research Sub-Questions", level=1)
    for i, q in enumerate(sub_questions, 1):
        doc.add_paragraph(f"{i}. {q}")

    # ─── 4. Findings per Sub-Question ────────────────────────
    doc.add_heading("Findings", level=1)
    for q, finding in research_findings.items():
        doc.add_heading(q, level=2)
        for para in finding.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # ─── 5. Coverage Gaps Analysis ───────────────────────────
    doc.add_heading("Coverage Gaps Analysis", level=1)
    if reflection_gaps:
        doc.add_paragraph("The following gaps were identified during research:")
        for gap in reflection_gaps:
            doc.add_paragraph(f"• {gap}")
        doc.add_paragraph("These gaps were addressed through additional targeted research.")
    else:
        doc.add_paragraph("No significant coverage gaps were identified. The research provides comprehensive coverage of the topic.")

    # ─── 6. Final Synthesis ──────────────────────────────────
    doc.add_heading("Final Synthesis", level=1)
    for para in synth_paragraphs:
        doc.add_paragraph(para)

    # ─── 7. References ───────────────────────────────────────
    doc.add_heading("References", level=1)
    ref_num = 1
    seen = set()
    for finding in research_findings.values():
        # Extract document references
        import re
        doc_refs = re.findall(r'\[([^\]]+\.pdf[^\]]*)\]', finding)
        for ref in doc_refs:
            if ref not in seen:
                seen.add(ref)
                doc.add_paragraph(f"[{ref_num}] {ref}")
                ref_num += 1
        # Extract web references
        web_refs = re.findall(r'\[Web: ([^\]]+)\]\(([^)]+)\)', finding)
        for title, url in web_refs:
            key = f"{title}:{url}"
            if key not in seen:
                seen.add(key)
                doc.add_paragraph(f"[{ref_num}] {title} — {url}")
                ref_num += 1

    if ref_num == 1:
        doc.add_paragraph("No external references cited.")

    if web_used:
        doc.add_paragraph()
        note = doc.add_paragraph()
        run = note.add_run("Note: This report includes information from web search results. "
                          "Web sources may change over time.")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(filepath)
